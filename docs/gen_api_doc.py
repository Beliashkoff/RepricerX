"""
Генератор docs/api_endpoints.docx — справочник по REST API RepricerX.
Запускать из корня репозитория: python3 docs/gen_api_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── палитра ──────────────────────────────────────────────────────────────────
COLOR_HEADER  = RGBColor(0x1F, 0x35, 0x64)   # тёмно-синий (заголовки)
COLOR_GET     = RGBColor(0x00, 0x7B, 0x00)   # зелёный
COLOR_POST    = RGBColor(0x00, 0x52, 0xCC)   # синий
COLOR_PATCH   = RGBColor(0xB3, 0x5B, 0x00)   # оранжевый
COLOR_DELETE  = RGBColor(0xCC, 0x00, 0x00)   # красный
COLOR_MONO_BG = RGBColor(0xF3, 0xF4, 0xF6)   # фон кода
COLOR_SECTION = RGBColor(0xE8, 0xED, 0xF5)   # фон заголовка секции

METHOD_COLOR = {
    "GET":    COLOR_GET,
    "POST":   COLOR_POST,
    "PATCH":  COLOR_PATCH,
    "DELETE": COLOR_DELETE,
}

# ── вспомогательные функции ───────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_mono(para, text, bold=False):
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.bold = bold
    return run

def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = COLOR_HEADER
    p.runs[0].font.size = Pt(16)

def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = COLOR_HEADER
    p.runs[0].font.size = Pt(13)

def section_label(doc, text):
    """Серый фон-строка как подзаголовок эндпоинта."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, "E8EDF5")
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
    doc.add_paragraph()   # отступ

def add_kv_table(doc, rows_data: list[tuple]):
    """Двухколоночная таблица ключ / значение."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)
    for key, val in rows_data:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(row.cells[0], "F3F4F6")
        p = row.cells[1].paragraphs[0]
        add_mono(p, val)
    doc.add_paragraph()


def add_endpoint(doc, method: str, path: str, summary: str,
                 auth_required: bool,
                 csrf_required: bool,
                 description: str = "",
                 request_fields: list[tuple] | None = None,
                 response_fields: list[tuple] | None = None,
                 error_codes: list[tuple] | None = None,
                 notes: list[str] | None = None):
    """Добавляет полное описание одного эндпоинта."""

    # ── заголовок метод + путь ────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r_method = p.add_run(f"{method}  ")
    r_method.bold = True
    r_method.font.size = Pt(12)
    r_method.font.color.rgb = METHOD_COLOR.get(method, COLOR_HEADER)
    r_path = p.add_run(path)
    r_path.font.name = "Courier New"
    r_path.font.size = Pt(12)
    r_path.bold = True

    # ── краткое описание ──────────────────────────────────────────────────
    p2 = doc.add_paragraph(summary)
    p2.runs[0].font.size = Pt(10)
    p2.runs[0].italic = True

    # ── мета ─────────────────────────────────────────────────────────────
    meta = []
    meta.append(("Авторизация", "Требуется (cookie сессии)" if auth_required else "Не требуется"))
    meta.append(("CSRF-защита",  "Да (Origin/Referer)" if csrf_required else "Нет"))
    add_kv_table(doc, meta)

    # ── описание ─────────────────────────────────────────────────────────
    if description:
        dp = doc.add_paragraph(description)
        dp.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # ── тело запроса ─────────────────────────────────────────────────────
    if request_fields:
        section_label(doc, "Тело запроса (application/json)")
        t = doc.add_table(rows=0, cols=3)
        t.style = "Table Grid"
        hdr = t.add_row()
        for i, h in enumerate(["Поле", "Тип", "Описание"]):
            hdr.cells[i].text = h
            hdr.cells[i].paragraphs[0].runs[0].bold = True
            hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(hdr.cells[i], "D9E2F3")
        for field, ftype, fdesc in request_fields:
            row = t.add_row()
            add_mono(row.cells[0].paragraphs[0], field, bold=True)
            add_mono(row.cells[1].paragraphs[0], ftype)
            row.cells[2].paragraphs[0].add_run(fdesc).font.size = Pt(9)
        doc.add_paragraph()

    # ── тело ответа ───────────────────────────────────────────────────────
    if response_fields:
        section_label(doc, "Тело ответа (application/json)")
        t = doc.add_table(rows=0, cols=3)
        t.style = "Table Grid"
        hdr = t.add_row()
        for i, h in enumerate(["Поле", "Тип", "Описание"]):
            hdr.cells[i].text = h
            hdr.cells[i].paragraphs[0].runs[0].bold = True
            hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(hdr.cells[i], "D9E2F3")
        for field, ftype, fdesc in response_fields:
            row = t.add_row()
            add_mono(row.cells[0].paragraphs[0], field, bold=True)
            add_mono(row.cells[1].paragraphs[0], ftype)
            row.cells[2].paragraphs[0].add_run(fdesc).font.size = Pt(9)
        doc.add_paragraph()

    # ── коды ошибок ───────────────────────────────────────────────────────
    if error_codes:
        section_label(doc, "Коды ошибок")
        t = doc.add_table(rows=0, cols=3)
        t.style = "Table Grid"
        hdr = t.add_row()
        for i, h in enumerate(["HTTP", "code", "Описание"]):
            hdr.cells[i].text = h
            hdr.cells[i].paragraphs[0].runs[0].bold = True
            hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(hdr.cells[i], "D9E2F3")
        for http_code, err_code, desc in error_codes:
            row = t.add_row()
            add_mono(row.cells[0].paragraphs[0], str(http_code))
            add_mono(row.cells[1].paragraphs[0], err_code, bold=True)
            row.cells[2].paragraphs[0].add_run(desc).font.size = Pt(9)
        doc.add_paragraph()

    # ── примечания ────────────────────────────────────────────────────────
    if notes:
        for note in notes:
            np = doc.add_paragraph(f"⚑  {note}", style="List Bullet")
            np.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    doc.add_paragraph("─" * 90).runs[0].font.size = Pt(7)


# ── формат ошибки ─────────────────────────────────────────────────────────────
ERROR_ENVELOPE = [
    ("error.code",    "string", "Машиночитаемый код ошибки"),
    ("error.message", "string", "Человекочитаемое сообщение на русском"),
]

# ═══════════════════════════════════════════════════════════════════════════════
#  ГЛАВНЫЙ БЛОК
# ═══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── стили страницы ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# ── титульный блок ────────────────────────────────────────────────────────────
title = doc.add_heading("RepricerX — Справочник REST API", 0)
title.runs[0].font.color.rgb = COLOR_HEADER
title.runs[0].font.size = Pt(20)

updated = doc.add_paragraph(
    f"Версия документа: актуально на {datetime.date.today().strftime('%d.%m.%Y')}\n"
    "Базовый URL: https://repricerx.ru/api  |  Протокол: HTTPS\n"
    "Аутентификация: HttpOnly-cookie  session_token  (Set-Cookie при логине)"
)
updated.runs[0].font.size = Pt(10)
doc.add_paragraph()

# ── 0. Общие соглашения ───────────────────────────────────────────────────────
heading1(doc, "0. Общие соглашения")

doc.add_paragraph(
    "Все эндпоинты возвращают JSON (Content-Type: application/json).\n"
    "При ошибке тело всегда имеет форму:"
)
add_kv_table(doc, [
    ('{"error":{"code":...', '"...","message":"..."}}'),
])

doc.add_paragraph(
    "Защищённые эндпоинты требуют валидного cookie session_token.\n"
    "Мутирующие защищённые эндпоинты (POST/PATCH/DELETE) дополнительно проверяют "
    "заголовок Origin или Referer (same-origin CSRF-защита)."
)
doc.add_paragraph()

# ── 1. Сервисные эндпоинты ────────────────────────────────────────────────────
heading1(doc, "1. Сервисные эндпоинты")

add_endpoint(
    doc, "GET", "/healthz",
    summary="Проверка живости сервиса (liveness probe)",
    auth_required=False, csrf_required=False,
    description="Всегда возвращает 200 OK пока процесс запущен. Используется Docker/K8s liveness probe.",
    response_fields=[("status", "string", 'Всегда "ok"')],
)

add_endpoint(
    doc, "GET", "/ready",
    summary="Проверка готовности сервиса (readiness probe)",
    auth_required=False, csrf_required=False,
    description=(
        "Проверяет доступность Postgres (ping) и Redis (ping). "
        "Возвращает 200 если оба доступны, 503 если хотя бы один недоступен."
    ),
    response_fields=[
        ("status", "string", '"ok" или "unavailable"'),
        ("db",     "string", '"ok" или сообщение об ошибке'),
        ("redis",  "string", '"ok" или сообщение об ошибке'),
    ],
    error_codes=[(503, "service_unavailable", "Postgres или Redis недоступны")],
)

# ── 2. Аутентификация ─────────────────────────────────────────────────────────
heading1(doc, "2. Аутентификация  (/api/auth)")

add_endpoint(
    doc, "POST", "/api/auth/register",
    summary="Регистрация нового пользователя",
    auth_required=False, csrf_required=False,
    description=(
        "Создаёт учётную запись. Пароль хешируется алгоритмом argon2id. "
        "После успешной регистрации на указанный email отправляется письмо "
        "с ссылкой для подтверждения адреса."
    ),
    request_fields=[
        ("email",       "string", "Адрес электронной почты (уникальный)"),
        ("password",    "string", "Пароль: 12–128 символов, минимум 1 буква и 1 цифра"),
        ("displayName", "string", "Отображаемое имя пользователя"),
    ],
    response_fields=[
        ("email", "string", "Email зарегистрированного пользователя"),
    ],
    error_codes=[
        (400, "bad_request",    "Неверный формат тела запроса"),
        (400, "invalid_email",  "Некорректный формат email"),
        (400, "weak_password",  "Пароль не соответствует требованиям"),
        (409, "email_taken",    "Пользователь с таким email уже существует"),
        (500, "internal_error", "Внутренняя ошибка сервера"),
    ],
    notes=[
        "HTTP 201 Created при успехе.",
        "Новый пользователь имеет статус pending до подтверждения email. "
        "Функциональность до верификации может быть ограничена.",
    ],
)

add_endpoint(
    doc, "POST", "/api/auth/login",
    summary="Вход в систему",
    auth_required=False, csrf_required=False,
    description=(
        "Проверяет email и пароль (argon2id, constant-time). "
        "При успехе создаёт сессию и устанавливает HttpOnly-cookie session_token. "
        "TTL сессии: 24 ч idle / 7 суток absolute."
    ),
    request_fields=[
        ("email",    "string", "Адрес электронной почты"),
        ("password", "string", "Пароль пользователя"),
    ],
    response_fields=[
        ("id",          "string (UUID)", "Идентификатор пользователя"),
        ("email",       "string",        "Email пользователя"),
        ("displayName", "string",        "Отображаемое имя"),
    ],
    error_codes=[
        (401, "invalid_credentials", "Неверный email или пароль (намеренно общий ответ)"),
        (403, "user_blocked",        "Аккаунт заблокирован"),
        (500, "internal_error",      "Внутренняя ошибка сервера"),
    ],
    notes=[
        "Set-Cookie: session_token=<token>; HttpOnly; SameSite=Strict; Path=/",
        "Неверный JSON возвращает 401 (не 400) — не раскрывает причину отказа.",
        "После N неудачных попыток аккаунт временно блокируется.",
    ],
)

add_endpoint(
    doc, "POST", "/api/auth/logout",
    summary="Выход из системы",
    auth_required=True, csrf_required=True,
    description=(
        "Инвалидирует текущую сессию в БД и удаляет cookie. "
        "Если cookie отсутствует или уже истёк — всё равно возвращает 204."
    ),
    response_fields=[],
    error_codes=[
        (401, "unauthorized", "Нет действующей сессии"),
    ],
    notes=[
        "HTTP 204 No Content при успехе.",
        "Требует заголовок Origin совпадающий с доменом сервиса (CSRF).",
    ],
)

add_endpoint(
    doc, "GET", "/api/auth/me",
    summary="Получение профиля текущего пользователя",
    auth_required=True, csrf_required=False,
    description="Возвращает данные аккаунта, привязанного к активной сессии.",
    response_fields=[
        ("id",          "string (UUID)",      "Идентификатор пользователя"),
        ("email",       "string",             "Email пользователя"),
        ("displayName", "string",             "Отображаемое имя"),
        ("status",      "string (enum)",      "Статус: active | pending | blocked"),
        ("createdAt",   "string (RFC 3339)",  "Дата регистрации (UTC)"),
    ],
    error_codes=[
        (401, "unauthorized", "Сессия отсутствует или истекла"),
        (403, "user_blocked", "Аккаунт заблокирован"),
    ],
)

add_endpoint(
    doc, "PATCH", "/api/auth/me",
    summary="Обновление отображаемого имени",
    auth_required=True, csrf_required=True,
    description="Изменяет displayName текущего пользователя.",
    request_fields=[
        ("displayName", "string", "Новое отображаемое имя (обязательное поле)"),
    ],
    response_fields=[
        ("id",          "string (UUID)", "Идентификатор пользователя"),
        ("email",       "string",        "Email пользователя"),
        ("displayName", "string",        "Обновлённое отображаемое имя"),
    ],
    error_codes=[
        (400, "bad_request",    "Отсутствует поле displayName"),
        (401, "unauthorized",   "Сессия отсутствует или истекла"),
        (500, "internal_error", "Ошибка обновления профиля"),
    ],
    notes=["Требует заголовок Origin совпадающий с доменом сервиса (CSRF)."],
)

add_endpoint(
    doc, "GET", "/api/auth/verify",
    summary="Подтверждение адреса электронной почты",
    auth_required=False, csrf_required=False,
    description=(
        "Принимает токен из письма верификации. "
        "При успехе или ошибке выполняет редирект на фронтенд "
        "c параметром ?verified=1 или ?verified=0."
    ),
    request_fields=[
        ("token", "string (query param)", "Одноразовый токен из ссылки в письме"),
    ],
    response_fields=[],
    error_codes=[],
    notes=[
        "HTTP 302 Found — всегда редирект, нет JSON-ответа.",
        "Успех: /login?verified=1 | Ошибка: /login?verified=0",
    ],
)

add_endpoint(
    doc, "POST", "/api/auth/verification/resend",
    summary="Повторная отправка письма верификации",
    auth_required=False, csrf_required=False,
    description=(
        "Инвалидирует предыдущий токен верификации и отправляет новое письмо. "
        "Всегда возвращает 202 — не раскрывает существование email."
    ),
    request_fields=[
        ("email", "string", "Email для повторной отправки"),
    ],
    response_fields=[],
    error_codes=[],
    notes=[
        "HTTP 202 Accepted при любом результате (anti-enumeration).",
    ],
)

# ── 3. Планируемые эндпоинты ──────────────────────────────────────────────────
heading1(doc, "3. Планируемые эндпоинты (в разработке)")

doc.add_paragraph(
    "Следующие группы эндпоинтов будут добавлены по мере реализации этапов плана разработки. "
    "Описание будет дополнено в этом документе."
)

planned = [
    ("Этап 2 — Магазины",          "/api/shops",            "CRUD магазинов, тест подключения к WB/Ozon"),
    ("Этап 3 — Каталог SKU",       "/api/products",         "Импорт, ручное добавление, поиск, редактирование"),
    ("Этап 4 — Стратегии",         "/api/strategies",       "CRUD стратегий, назначение на SKU"),
    ("Этап 5 — Расчёт цен",        "/api/pricing",          "Симуляция и формирование плана изменений"),
    ("Этап 6 — Отправка цен",      "/api/pricing/dispatch", "Утверждение и отправка плана в маркетплейс"),
    ("Этап 7 — Планировщик",       "/api/shops/:id/run-now","Ручной триггер пересчёта"),
    ("Этап 8 — Журнал и отчёты",   "/api/audit",            "Журнал изменений цен, экспорт CSV, метрики"),
]

t = doc.add_table(rows=0, cols=3)
t.style = "Table Grid"
hdr = t.add_row()
for i, h in enumerate(["Этап", "Префикс", "Описание"]):
    hdr.cells[i].text = h
    hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(hdr.cells[i], "D9E2F3")
for stage, prefix, desc in planned:
    row = t.add_row()
    row.cells[0].paragraphs[0].add_run(stage).font.size = Pt(9)
    add_mono(row.cells[1].paragraphs[0], prefix)
    row.cells[2].paragraphs[0].add_run(desc).font.size = Pt(9)

doc.add_paragraph()

# ── футер ─────────────────────────────────────────────────────────────────────
footer_p = doc.add_paragraph(
    f"Документ сгенерирован автоматически из исходного кода · {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}"
)
footer_p.runs[0].font.size = Pt(8)
footer_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── сохранение ────────────────────────────────────────────────────────────────
import os
out_path = os.path.join(os.path.dirname(__file__), "api_endpoints.docx")
doc.save(out_path)
print(f"Сохранено: {out_path}")
